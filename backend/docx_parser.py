"""Word (.docx) document parser.

Produces a ``DocData`` object whose ``sections`` are the *leaf* sections of the
heading hierarchy (deepest headings with no child headings beneath them) and
whose ``fig_table_dict`` lists all figure/table captions with their in-text
mentions.
"""

from __future__ import annotations

import re
from dataclasses import dataclass
from typing import Any

from docx import Document  # python-docx

from doc_models import (
    CAPTION_RE as _CAPTION_RE,
    MENTION_RE as _MENTION_RE,
    DocData,
    FigTableEntry,
    FigTableMention,
    Section,
    normalise_mention as _normalise_mention,
)

_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


# ---------------------------------------------------------------------------
# Public entry point
# ---------------------------------------------------------------------------

def parse_docx(file_path: str, range_spec: dict | None) -> DocData:
    """Parse *file_path* (.docx) into a :class:`DocData`."""
    from field_updater import update_docx_fields, cleanup_updated_docx
    updated_path = update_docx_fields(file_path)
    try:
        return _parse_docx_inner(updated_path, file_path, range_spec)
    finally:
        cleanup_updated_docx(file_path, updated_path)


# ---------------------------------------------------------------------------
# Internal helpers — heading detection
# ---------------------------------------------------------------------------

@dataclass
class _HeadingInfo:
    level: int
    number: str
    title: str
    para_idx: int  # index in the flat paragraph list


def _heading_level(para) -> int:
    """Return heading level (1-based) or 0 if not a heading."""
    style = para.style.name.lower()
    for prefix in ("heading ", "заголовок "):
        if style.startswith(prefix):
            try:
                return int(style[len(prefix):].strip().split()[0])
            except (ValueError, IndexError):
                return 1
    return 0


def _extract_heading_number(title: str) -> str:
    """Try to pull a leading numbering token from a heading title.

    E.g. "3.2 Методы исследования" → "3.2".
    Returns "" if no numeric prefix is found (auto-numbering will be applied later).
    """
    m = re.match(r"^([\d]+(?:[.\-–—][\d]+)*)\s*", title)
    if m:
        return m.group(1)
    return ""


def _apply_auto_numbering(headings: list[_HeadingInfo]) -> None:
    """Assign synthetic section numbers to headings that have no explicit numeric prefix.

    Only applied when ALL headings lack explicit numbers (i.e. the document uses
    Word automatic list-numbering where the number is not part of para.text).
    """
    has_explicit = any(h.number for h in headings)
    if has_explicit:
        return

    counters: dict[int, int] = {}
    for h in headings:
        level = h.level
        counters[level] = counters.get(level, 0) + 1
        for deeper in [k for k in counters if k > level]:
            del counters[deeper]
        parts = [str(counters[lvl]) for lvl in sorted(counters)]
        h.number = ".".join(parts)


# ---------------------------------------------------------------------------
# Internal helpers — XML field resolution
# ---------------------------------------------------------------------------

def _styleref_level(instr: str) -> int:
    """Extract heading level from a STYLEREF instruction.

    Examples:
        ' STYLEREF "Заголовок 1" \\s ' → 1
        ' STYLEREF "Heading 2" \\s '   → 2
    """
    m = re.search(r'(?:heading|заголовок)[\s\xa0]+(\d+)', instr, re.IGNORECASE)
    if m:
        return int(m.group(1))
    m2 = re.search(r'STYLEREF\s+"(\d+)"', instr, re.IGNORECASE)
    return int(m2.group(1)) if m2 else 1


def _resolve_para_fields(
    para_e: Any,
    *,
    section_number: str = "",
    bm_to_label: dict[str, str] | None = None,
) -> str:
    """Resolve Word field codes in *para_e* and return the paragraph text.

    - STYLEREF fields are replaced with the appropriate prefix of *section_number*
      (only when *section_number* is non-empty).
    - REF fields are replaced with the resolved label from *bm_to_label*
      (only when *bm_to_label* is provided and the bookmark is found).
    - All other field results (e.g. SEQ counters) are kept verbatim.
    - Non-breaking hyphens (``<w:noBreakHyphen>``) are emitted as U+2011
      unless they fall inside a field result that will be replaced.
    """
    parts: list[str] = []
    in_field = False
    in_field_result = False
    is_styleref = False
    is_ref = False
    ref_label: str = ""
    level = 1

    for node in para_e.iter():
        tag = node.tag
        if tag == f"{{{_W}}}fldChar":
            fld_type = node.get(f"{{{_W}}}fldCharType", "")
            if fld_type == "begin":
                in_field = True
                in_field_result = False
                is_styleref = False
                is_ref = False
                ref_label = ""
                level = 1
            elif fld_type == "separate":
                in_field_result = True
            elif fld_type == "end":
                if is_styleref and section_number:
                    num_parts = section_number.split(".")
                    parts.append(".".join(num_parts[:level]))
                elif is_ref and ref_label:
                    parts.append(ref_label)
                in_field = False
                in_field_result = False
                is_styleref = False
                is_ref = False
                ref_label = ""
        elif tag == f"{{{_W}}}instrText":
            instr = (node.text or "").strip()
            if section_number and re.match(r"STYLEREF", instr, re.IGNORECASE):
                is_styleref = True
                level = _styleref_level(instr)
            elif bm_to_label is not None and instr.upper().startswith("REF "):
                instr_parts = instr.split()
                if len(instr_parts) >= 2:
                    bm = instr_parts[1]
                    if bm in bm_to_label:
                        is_ref = True
                        ref_label = bm_to_label[bm]
        elif tag == f"{{{_W}}}noBreakHyphen":
            if not in_field or (in_field_result and not is_styleref and not is_ref):
                parts.append("\u2011")
        elif tag == f"{{{_W}}}t":
            if not in_field:
                parts.append(node.text or "")
            elif in_field_result and not is_styleref and not is_ref:
                parts.append(node.text or "")

    return "".join(parts)


def _para_text_elem(para_elem: Any) -> str:
    return "".join(node.text or "" for node in para_elem.iter(f"{{{_W}}}t"))


# ---------------------------------------------------------------------------
# Internal helpers — section range filtering
# ---------------------------------------------------------------------------

def _section_in_range(number: str, items: list[dict]) -> bool:
    """Return True if *number* falls within any of the given section ranges."""
    def _key(s: str) -> tuple[int, ...]:
        result = []
        for p in re.split(r"[.\-–—]", s):
            try:
                result.append(int(p))
            except ValueError:
                pass
        return tuple(result)

    num_key = _key(number)
    if not num_key:
        return False
    for item in items:
        start_key = _key(item.get("start", ""))
        end_key = _key(item.get("end", item.get("start", "")))
        if not start_key:
            continue
        if start_key <= num_key <= end_key:
            return True
        prefix = num_key[:len(start_key)]
        if prefix and start_key <= prefix <= end_key:
            return True
    return False


# ---------------------------------------------------------------------------
# Internal helpers — section list builder
# ---------------------------------------------------------------------------

def _build_elem_section_list(
    paras_elems: list,
    flat_paras,
    para_to_section: dict[int, str],
) -> list[str]:
    """Return a section-number string for every element in *paras_elems*.

    Table-cell paragraphs (absent from *flat_paras*) inherit the last seen
    top-level paragraph's section number.
    """
    flat_id_to_sec: dict[int, str] = {
        id(p._p): para_to_section.get(i, "")
        for i, p in enumerate(flat_paras)
    }
    result: list[str] = []
    current = ""
    for para_e in paras_elems:
        sec = flat_id_to_sec.get(id(para_e))
        if sec is not None:
            current = sec
        result.append(current)
    return result


# ---------------------------------------------------------------------------
# Figure/table dictionary builders
# ---------------------------------------------------------------------------

def _renumber_entries_by_section(
    label_order: list[str],
    entries: dict[str, FigTableEntry],
    bm_to_label: dict[str, str],
) -> tuple[dict[str, FigTableEntry], dict[str, str]]:
    """Renumber figure/table entries 1, 2, 3… within each (section, kind) group.

    Word's SEQ field uses a chapter-wide running counter; this reassigns
    sequential per-section numbers so the first figure in 5.3 gets label
    'Рисунок 5.3‑1', the second 'Рисунок 5.3‑2', etc.
    """
    sec_kind_count: dict[tuple[str, str], int] = {}
    remap: dict[str, str] = {}

    for old_label in label_order:
        entry = entries.get(old_label)
        if entry is None:
            continue
        m = re.match(r'^(Таблица|Рисунок)\b', old_label, re.IGNORECASE)
        if not m:
            remap[old_label] = old_label
            continue
        kind = m.group(1)
        key = (entry.section_number, kind)
        sec_kind_count[key] = sec_kind_count.get(key, 0) + 1
        remap[old_label] = f"{kind} {entry.section_number}\u2011{sec_kind_count[key]}"

    new_entries: dict[str, FigTableEntry] = {}
    for old_label, entry in entries.items():
        new_label = remap.get(old_label, old_label)
        new_caption = (
            entry.caption.replace(old_label, new_label, 1)
            if old_label in entry.caption
            else entry.caption
        )
        new_entries[new_label] = FigTableEntry(
            label=new_label,
            caption=new_caption,
            section_number=entry.section_number,
            mentions=entry.mentions,
        )

    new_bm_to_label = {bm: remap.get(lbl, lbl) for bm, lbl in bm_to_label.items()}
    return new_entries, new_bm_to_label


def _build_fig_table_dict_docx(
    doc: Document,
    sections: list[Section],
    para_to_section: dict[int, str] | None = None,
) -> list[FigTableEntry]:
    """Build a list of FigTableEntry from the whole document.

    Uses bookmark-based REF fields for docx plus a fallback text scan.

    *para_to_section* maps flat-paragraph index → section number.  When
    provided it is used directly; otherwise it is rebuilt here via
    _extract_heading_number (legacy path).
    """
    body = doc.element.body
    paras_elems = list(body.iter(f"{{{_W}}}p"))
    flat_paras = doc.paragraphs

    if para_to_section is None:
        para_to_section = {}
        current_sec = ""
        for idx, para in enumerate(flat_paras):
            if _heading_level(para) > 0:
                current_sec = _extract_heading_number(para.text.strip())
            para_to_section[idx] = current_sec

    elem_section = _build_elem_section_list(paras_elems, flat_paras, para_to_section)

    entries: dict[str, FigTableEntry] = {}
    bm_to_label: dict[str, str] = {}

    for para_idx_e, para_e in enumerate(paras_elems):
        sec_num = elem_section[para_idx_e]
        text = _resolve_para_fields(para_e, section_number=sec_num).strip()
        cap_m = _CAPTION_RE.match(text)
        if cap_m:
            label = cap_m.group(0).strip()
            entry = FigTableEntry(label=label, caption=text, section_number=sec_num)
            entries[label] = entry
            for bm in para_e.iter(f"{{{_W}}}bookmarkStart"):
                bm_name = bm.get(f"{{{_W}}}name", "")
                if bm_name:
                    bm_to_label[bm_name] = label

    entries, bm_to_label = _renumber_entries_by_section(
        list(entries.keys()), entries, bm_to_label
    )

    # Collect mentions via REF fields
    for para_idx_e, para_e in enumerate(paras_elems):
        instr_texts = [node.text or "" for node in para_e.iter(f"{{{_W}}}instrText")]
        for instr in instr_texts:
            instr = instr.strip()
            if not instr.upper().startswith("REF "):
                continue
            parts = instr.split()
            if len(parts) < 2:
                continue
            bm_name = parts[1]
            if bm_name not in bm_to_label:
                continue
            label = bm_to_label[bm_name]
            ctx = []
            if para_idx_e > 0:
                ctx.append(_resolve_para_fields(paras_elems[para_idx_e - 1], bm_to_label=bm_to_label).strip())
            ctx.append(_resolve_para_fields(para_e, bm_to_label=bm_to_label).strip())
            if para_idx_e < len(paras_elems) - 1:
                ctx.append(_resolve_para_fields(paras_elems[para_idx_e + 1], bm_to_label=bm_to_label).strip())
            context = "\n".join(p for p in ctx if p)
            sec_num = elem_section[para_idx_e]
            entries[label].mentions.append(FigTableMention(context=context, section_number=sec_num))

    # Fallback: text scan for mentions not covered by REF fields
    mentioned_via_ref: set[str] = set()
    for entry in entries.values():
        for mention in entry.mentions:
            mentioned_via_ref.add(mention.context[:50])

    for para_idx_e, para_e in enumerate(paras_elems):
        text = _para_text_elem(para_e).strip()
        for m in _MENTION_RE.finditer(text):
            label_candidate = _normalise_mention(m.group(0))
            if label_candidate not in entries:
                continue
            ctx = []
            if para_idx_e > 0:
                ctx.append(_para_text_elem(paras_elems[para_idx_e - 1]).strip())
            ctx.append(text)
            if para_idx_e < len(paras_elems) - 1:
                ctx.append(_para_text_elem(paras_elems[para_idx_e + 1]).strip())
            context = "\n".join(p for p in ctx if p)
            if context[:50] in mentioned_via_ref:
                continue
            sec_num = elem_section[para_idx_e]
            entries[label_candidate].mentions.append(
                FigTableMention(context=context, section_number=sec_num)
            )
            mentioned_via_ref.add(context[:50])

    return list(entries.values())


# ---------------------------------------------------------------------------
# Main DOCX parsing pass
# ---------------------------------------------------------------------------

def _parse_docx_inner(doc_path: str, original_path: str, range_spec: dict | None) -> DocData:
    doc = Document(doc_path)
    paras = doc.paragraphs

    # Pass 1: collect and number headings
    headings: list[_HeadingInfo] = []
    for idx, para in enumerate(paras):
        level = _heading_level(para)
        if level == 0:
            continue
        title = para.text.strip()
        number = _extract_heading_number(title)
        headings.append(_HeadingInfo(level=level, number=number, title=title, para_idx=idx))

    _apply_auto_numbering(headings)

    # Build para_idx → section_number map
    heading_num_by_para = {h.para_idx: h.number for h in headings}
    p2s: dict[int, str] = {}
    cur_sec = ""
    for idx in range(len(paras)):
        if idx in heading_num_by_para:
            cur_sec = heading_num_by_para[idx]
        p2s[idx] = cur_sec

    # Pass 2: identify leaf headings
    leaf_indices: set[int] = set()
    for i, h in enumerate(headings):
        is_leaf = True
        for j in range(i + 1, len(headings)):
            nxt = headings[j]
            if nxt.level <= h.level:
                break
            if nxt.level > h.level:
                is_leaf = False
                break
        if is_leaf:
            leaf_indices.add(i)

    # No headings: treat whole document as one section
    if not headings:
        full_text = "\n".join(p.text.strip() for p in paras if p.text.strip())
        early_sections = [Section(number="", title="Документ", text=full_text, level=0)]
        fig_table_dict = _build_fig_table_dict_docx(doc, early_sections, para_to_section=None)
        return DocData(fmt="docx", file_path=original_path,
                       sections=early_sections, fig_table_dict=fig_table_dict, raw_docx=doc)

    # Preliminary: build bm_to_label for REF-field resolution in body text
    paras_elems_pre = list(doc.element.body.iter(f"{{{_W}}}p"))
    elem_sec_pre = _build_elem_section_list(paras_elems_pre, paras, p2s)

    pre_raw: list[tuple[str, str, list[str]]] = []  # (raw_label, sec, [bm_names])
    for pe_idx, pe in enumerate(paras_elems_pre):
        cap_text = _resolve_para_fields(pe, section_number=elem_sec_pre[pe_idx]).strip()
        cap_m = _CAPTION_RE.match(cap_text)
        if cap_m:
            lbl = cap_m.group(0).strip()
            bm_names = [
                bm.get(f"{{{_W}}}name", "")
                for bm in pe.iter(f"{{{_W}}}bookmarkStart")
                if bm.get(f"{{{_W}}}name", "")
            ]
            pre_raw.append((lbl, elem_sec_pre[pe_idx], bm_names))

    # Renumber unique labels per (section, kind)
    pre_sk: dict[tuple[str, str], int] = {}
    pre_remap: dict[str, str] = {}
    for lbl, sec, _ in pre_raw:
        if lbl in pre_remap:
            continue
        km = re.match(r'^(Таблица|Рисунок)\b', lbl, re.IGNORECASE)
        if not km:
            continue
        kind = km.group(1)
        key = (sec, kind)
        pre_sk[key] = pre_sk.get(key, 0) + 1
        pre_remap[lbl] = f"{kind} {sec}\u2011{pre_sk[key]}"

    bm_to_label: dict[str, str] = {}
    for lbl, _sec, bm_names in pre_raw:
        new_lbl = pre_remap.get(lbl, lbl)
        for bm_name in bm_names:
            bm_to_label[bm_name] = new_lbl

    # Pass 3: collect text for each leaf section
    sections: list[Section] = []
    for i, h in enumerate(headings):
        if i not in leaf_indices:
            continue
        end_idx = headings[i + 1].para_idx if i + 1 < len(headings) else len(paras)

        text_parts = [h.title]
        for idx in range(h.para_idx + 1, end_idx):
            t = _resolve_para_fields(
                paras[idx]._p,
                section_number=p2s.get(idx, h.number),
                bm_to_label=bm_to_label,
            ).strip()
            if t:
                text_parts.append(t)

        sec_text = "\n".join(text_parts)
        for old, new in pre_remap.items():
            if old in sec_text:
                sec_text = sec_text.replace(old, new)

        sections.append(Section(number=h.number, title=h.title, text=sec_text, level=h.level))

    # Filter by range_spec
    if range_spec and range_spec.get("type") == "sections":
        items = range_spec.get("items", [])
        if items:
            sections = [s for s in sections if _section_in_range(s.number, items)]

    # Build figure/table dictionary
    fig_table_dict = _build_fig_table_dict_docx(doc, sections, para_to_section=p2s)

    if range_spec and range_spec.get("type") == "sections":
        items = range_spec.get("items", [])
        if items:
            fig_table_dict = [
                e for e in fig_table_dict
                if _section_in_range(e.section_number, items)
            ]

    return DocData(fmt="docx", file_path=original_path,
                   sections=sections, fig_table_dict=fig_table_dict, raw_docx=doc)
