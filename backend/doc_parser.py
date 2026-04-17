"""Unified document parser for Word (.docx) and PDF files.

Produces a ``DocData`` object that checkpoints consume regardless of the
original file format.  For docx, text is split into *leaf sections* (the
deepest heading level that has no child headings beneath it).  For pdf, each
page becomes one section.  The parser also pre-builds a dictionary of all
table/figure captions and the passages that reference them.
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Optional

import fitz  # PyMuPDF — used only for PDF
from docx import Document  # python-docx — used only for .docx

_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

# U+2011 NON-BREAKING HYPHEN is used by Word between chapter and figure numbers
_DASH = r"[-–—‑]"

_CAPTION_RE = re.compile(
    r"^(Таблица|Рисунок)\s+([\d\.]+[-–—‑][\d\.]+|[\d\.]+)(?=\s*[-–—‑]|\s*$)",
    re.IGNORECASE,
)

# Patterns for in-text references to figures/tables (e.g. "см. таблицу 3.1-2")
_MENTION_RE = re.compile(
    r"(таблиц[еуией]?\s+[\d\.]+[-–—‑][\d\.]+|рисун[окке]+\s+[\d\.]+[-–—‑][\d\.]+|"
    r"таблиц[еуией]?\s+[\d\.]+|рисун[окке]+\s+[\d\.]+)",
    re.IGNORECASE,
)


@dataclass
class Section:
    number: str   # "3.2" / "1.1.2" / "Страница 5"
    title: str    # заголовок раздела или ""
    text: str     # весь текст раздела
    level: int    # уровень заголовка (1-based); 0 для страниц PDF


@dataclass
class FigTableMention:
    context: str          # абзац(ы) где встречается ссылка
    section_number: str   # номер раздела где встречается ссылка


@dataclass
class FigTableEntry:
    label: str                       # "Таблица 3.1-2"
    caption: str                     # полный текст подписи
    section_number: str              # раздел где находится сама подпись
    mentions: list[FigTableMention] = field(default_factory=list)


@dataclass
class DocData:
    fmt: str                                    # "docx" | "pdf"
    file_path: str
    sections: list[Section] = field(default_factory=list)
    fig_table_dict: list[FigTableEntry] = field(default_factory=list)
    raw_docx: Optional[Document] = field(default=None, repr=False)
    raw_pdf: Optional[fitz.Document] = field(default=None, repr=False)


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def parse_document(
    file_path: str,
    range_spec: dict | None = None,
) -> DocData:
    """Detect format and parse *file_path* into a :class:`DocData`.

    *range_spec* is a dict produced by AI range validation::

        {"type": "sections"|"pages",
         "items": [{"start": "3.1", "end": "3.4"}, ...]}

    If *range_spec* is None the full document is processed.
    """
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"Файл не найден: {file_path}")

    ext = path.suffix.lower()
    if ext == ".docx":
        return _parse_docx(file_path, range_spec)
    elif ext == ".pdf":
        return _parse_pdf(file_path, range_spec)
    else:
        raise ValueError(f"Неподдерживаемый формат файла: {ext}. Ожидается .docx или .pdf")


# ---------------------------------------------------------------------------
# Range helpers
# ---------------------------------------------------------------------------

def _section_in_range(number: str, items: list[dict]) -> bool:
    """Return True if *number* falls within any of the given ranges.

    Comparison is done as a dot-separated tuple of ints (best-effort).
    """
    def _key(s: str) -> tuple[int, ...]:
        parts = re.split(r"[.\-–—]", s)
        result = []
        for p in parts:
            try:
                result.append(int(p))
            except ValueError:
                pass
        return tuple(result)

    num_key = _key(number)
    if not num_key:
        return False
    for item in items:
        start_key = _key(item.get("start", ""))
        end_key = _key(item.get("end", item.get("start", "")))
        if not start_key:
            continue
        # Direct range check (same-level sections)
        if start_key <= num_key <= end_key:
            return True
        # Subsection/prefix check: "5.1", "5.1.2" should match range "5"–"5"
        # Truncate num_key to the length of start_key and compare the prefix
        prefix = num_key[:len(start_key)]
        if prefix and start_key <= prefix <= end_key:
            return True
    return False


def _page_in_range(page_num: int, items: list[dict]) -> bool:
    for item in items:
        try:
            s = int(re.sub(r"\D", "", str(item.get("start", page_num))))
            e = int(re.sub(r"\D", "", str(item.get("end", item.get("start", page_num)))))
            if s <= page_num <= e:
                return True
        except (ValueError, TypeError):
            pass
    return False


# ---------------------------------------------------------------------------
# Word parser
# ---------------------------------------------------------------------------

def _heading_level(para) -> int:
    """Return heading level (1-based) or 0 if not a heading."""
    style = para.style.name.lower()
    for prefix in ("heading ", "заголовок "):
        if style.startswith(prefix):
            try:
                return int(style[len(prefix):].strip().split()[0])
            except (ValueError, IndexError):
                return 1
    return 0


def _extract_heading_number(title: str) -> str:
    """Try to pull a leading numbering token from a heading title.

    E.g. "3.2 Методы исследования" → "3.2".
    Returns "" if no numeric prefix is found (auto-numbering will be applied later).
    """
    m = re.match(r"^([\d]+(?:[.\-–—][\d]+)*)\s*", title)
    if m:
        return m.group(1)
    return ""


def _apply_auto_numbering(headings: list["_HeadingInfo"]) -> None:
    """Assign synthetic section numbers to headings that have no explicit numeric prefix.

    Uses the heading hierarchy order to generate "1", "1.1", "2.3" etc.
    Only applied when ALL headings lack explicit numbers (i.e. the document uses
    Word automatic list-numbering where the number is not part of para.text).
    """
    # If any heading already has an explicit number, keep the existing behaviour
    # (mixed documents with some numbered, some not).
    has_explicit = any(h.number for h in headings)
    if has_explicit:
        return

    counters: dict[int, int] = {}
    for h in headings:
        level = h.level
        counters[level] = counters.get(level, 0) + 1
        # Reset all deeper levels when ascending or staying at same level
        for deeper in [k for k in counters if k > level]:
            del counters[deeper]
        parts = [str(counters[lvl]) for lvl in sorted(counters)]
        h.number = ".".join(parts)


@dataclass
class _HeadingInfo:
    level: int
    number: str
    title: str
    para_idx: int  # index in the flat paragraph list


def _parse_docx(file_path: str, range_spec: dict | None) -> DocData:
    from field_updater import update_docx_fields, cleanup_updated_docx
    updated_path = update_docx_fields(file_path)
    try:
        return _parse_docx_inner(updated_path, file_path, range_spec)
    finally:
        cleanup_updated_docx(file_path, updated_path)


def _parse_docx_inner(doc_path: str, original_path: str, range_spec: dict | None) -> DocData:
    doc = Document(doc_path)
    paras = doc.paragraphs

    # --- Pass 1: collect all headings ---
    headings: list[_HeadingInfo] = []
    for idx, para in enumerate(paras):
        level = _heading_level(para)
        if level == 0:
            continue
        title = para.text.strip()
        number = _extract_heading_number(title)
        headings.append(_HeadingInfo(level=level, number=number, title=title, para_idx=idx))

    _apply_auto_numbering(headings)

    # Build para_idx → section_number using the already-numbered headings so
    # that _build_fig_table_dict_docx gets correct numbers (incl. auto-numbered).
    _heading_num_by_para = {h.para_idx: h.number for h in headings}
    _p2s: dict[int, str] = {}
    _cur_sec = ""
    for _idx in range(len(paras)):
        if _idx in _heading_num_by_para:
            _cur_sec = _heading_num_by_para[_idx]
        _p2s[_idx] = _cur_sec

    # --- Pass 2: identify leaf headings (no child heading before next same/parent) ---
    leaf_indices: set[int] = set()
    for i, h in enumerate(headings):
        is_leaf = True
        for j in range(i + 1, len(headings)):
            nxt = headings[j]
            if nxt.level <= h.level:
                break
            if nxt.level > h.level:
                is_leaf = False
                break
        if is_leaf:
            leaf_indices.add(i)

    # If document has no headings at all, treat the whole text as one section
    if not headings:
        full_text = "\n".join(p.text.strip() for p in paras if p.text.strip())
        sections = [Section(number="", title="Документ", text=full_text, level=0)]
        fig_table_dict = _build_fig_table_dict_docx(doc, sections, para_to_section=None)
        return DocData(fmt="docx", file_path=original_path,
                       sections=sections, fig_table_dict=fig_table_dict, raw_docx=doc)

    # --- Preliminary: build bm_to_label so REF fields in body text can be resolved ---
    _paras_elems_pre = list(doc.element.body.iter(f"{{{_W}}}p"))
    _elem_sec_pre = _build_elem_section_list(_paras_elems_pre, paras, _p2s)

    # Collect raw caption labels in document order, then renumber per-section
    _pre_raw: list[tuple[str, str, list[str]]] = []  # (raw_label, sec, [bm_names])
    _pre_seen: set[str] = set()  # dedup: same raw_label only counted once
    for _pe_idx, _pe in enumerate(_paras_elems_pre):
        _cap_text = _para_caption_text(_pe, _elem_sec_pre[_pe_idx]).strip()
        _cap_m = _CAPTION_RE.match(_cap_text)
        if _cap_m:
            _lbl = _cap_m.group(0).strip()
            _bm_names = [
                _bm.get(f"{{{_W}}}name", "")
                for _bm in _pe.iter(f"{{{_W}}}bookmarkStart")
                if _bm.get(f"{{{_W}}}name", "")
            ]
            _pre_raw.append((_lbl, _elem_sec_pre[_pe_idx], _bm_names))

    # Renumber unique labels per (section, kind)
    _pre_sk: dict[tuple[str, str], int] = {}
    _pre_remap: dict[str, str] = {}
    for _lbl, _sec, _ in _pre_raw:
        if _lbl in _pre_remap:
            continue  # already renumbered
        _km = re.match(r'^(Таблица|Рисунок)\b', _lbl, re.IGNORECASE)
        if not _km:
            continue
        _kind = _km.group(1)
        _key = (_sec, _kind)
        _pre_sk[_key] = _pre_sk.get(_key, 0) + 1
        _pre_remap[_lbl] = f"{_kind} {_sec}\u2011{_pre_sk[_key]}"

    _bm_to_label: dict[str, str] = {}
    for _lbl, _sec, _bm_names in _pre_raw:
        _new_lbl = _pre_remap.get(_lbl, _lbl)
        for _bm_name in _bm_names:
            _bm_to_label[_bm_name] = _new_lbl

    # --- Pass 3: collect text for each leaf section ---
    sections: list[Section] = []

    for i, h in enumerate(headings):
        if i not in leaf_indices:
            continue

        # Text spans from h.para_idx to the next heading's para_idx (exclusive)
        if i + 1 < len(headings):
            end_idx = headings[i + 1].para_idx
        else:
            end_idx = len(paras)

        text_parts = [h.title]
        for idx in range(h.para_idx + 1, end_idx):
            t = _para_full_resolved_text(paras[idx]._p, _p2s.get(idx, h.number), _bm_to_label).strip()
            if t:
                text_parts.append(t)

        sec_text = "\n".join(text_parts)
        # Caption paragraphs in body text still carry stale SEQ values (e.g. "5.3‑26").
        # Apply the same per-section renumbering map so section.text matches fig_table_dict.
        for _old, _new in _pre_remap.items():
            if _old in sec_text:
                sec_text = sec_text.replace(_old, _new)

        sections.append(Section(
            number=h.number,
            title=h.title,
            text=sec_text,
            level=h.level,
        ))

    # --- Filter by range_spec ---
    if range_spec and range_spec.get("type") == "sections":
        items = range_spec.get("items", [])
        if items:
            sections = [s for s in sections if _section_in_range(s.number, items)]

    # --- Build figure/table dictionary (always from whole doc, then filter) ---
    fig_table_dict = _build_fig_table_dict_docx(doc, sections, para_to_section=_p2s)

    if range_spec and range_spec.get("type") == "sections":
        items = range_spec.get("items", [])
        if items:
            fig_table_dict = [
                e for e in fig_table_dict
                if _section_in_range(e.section_number, items)
            ]

    return DocData(fmt="docx", file_path=original_path,
                   sections=sections, fig_table_dict=fig_table_dict, raw_docx=doc)


def _para_text_elem(para_elem: Any) -> str:
    return "".join(node.text or "" for node in para_elem.iter(f"{{{_W}}}t"))


def _styleref_level(instr: str) -> int:
    """Extract heading level from a STYLEREF instruction.

    Examples:
        ' STYLEREF "Заголовок 1" \\s ' → 1
        ' STYLEREF "Heading 2" \\s '   → 2
    """
    # "Heading N" / "Заголовок N" (standard, incl. non-breaking space)
    m = re.search(r'(?:heading|заголовок)[\s\xa0]+(\d+)', instr, re.IGNORECASE)
    if m:
        level = int(m.group(1))
    else:
        # Numeric style name: STYLEREF "N" (common in Russian Word templates)
        m2 = re.search(r'STYLEREF\s+"(\d+)"', instr, re.IGNORECASE)
        level = int(m2.group(1)) if m2 else 1
    return level


def _para_caption_text(para_e: Any, section_number: str) -> str:
    """Read paragraph text, substituting STYLEREF field results with the correct
    chapter-level number derived from *section_number*.

    Word caches the full heading title (e.g. 'Методы исследования...') as the
    STYLEREF result rather than the auto-list number.  We replace it with the
    appropriate prefix of *section_number* so a caption like
    'Рисунок Методы...-1' becomes 'Рисунок 5-1'.

    Non-STYLEREF fields (e.g. SEQ counters) are left untouched — their cached
    numeric result is used as-is.
    """
    parts: list[str] = []
    in_field = False
    in_field_result = False
    is_styleref = False
    level = 1

    for node in para_e.iter():
        tag = node.tag
        if tag == f"{{{_W}}}fldChar":
            fld_type = node.get(f"{{{_W}}}fldCharType", "")
            if fld_type == "begin":
                in_field = True
                in_field_result = False
                is_styleref = False
                level = 1
            elif fld_type == "separate":
                in_field_result = True
            elif fld_type == "end":
                if is_styleref and section_number:
                    # Use the first *level* components of the section number
                    num_parts = section_number.split(".")
                    parts.append(".".join(num_parts[:level]))
                in_field = False
                in_field_result = False
                is_styleref = False
        elif tag == f"{{{_W}}}instrText":
            instr = (node.text or "").strip()
            if re.match(r"STYLEREF", instr, re.IGNORECASE):
                is_styleref = True
                level = _styleref_level(instr)
        elif tag == f"{{{_W}}}noBreakHyphen":
            # Non-breaking hyphen stored as XML element (not <w:t>).
            # Emit it when outside a STYLEREF cached result.
            if not in_field or (in_field_result and not is_styleref):
                parts.append("\u2011")
        elif tag == f"{{{_W}}}t":
            if not in_field:
                parts.append(node.text or "")
            elif in_field_result and not is_styleref:
                # Non-STYLEREF field result (e.g. SEQ counter) — keep as-is
                parts.append(node.text or "")
            # STYLEREF result: skip the stale cached heading text

    return "".join(parts)


def _para_full_resolved_text(para_e: Any, section_number: str, bm_to_label: dict[str, str]) -> str:
    """Resolve both STYLEREF (using section_number) and REF (using bm_to_label) fields.

    Used when building section.text so that body paragraphs containing REF
    field references to captions show the resolved label instead of stale
    cached heading text.
    """
    parts: list[str] = []
    in_field = False
    in_field_result = False
    is_styleref = False
    is_ref = False
    ref_label: str = ""
    level = 1

    for node in para_e.iter():
        tag = node.tag
        if tag == f"{{{_W}}}fldChar":
            fld_type = node.get(f"{{{_W}}}fldCharType", "")
            if fld_type == "begin":
                in_field = True
                in_field_result = False
                is_styleref = False
                is_ref = False
                ref_label = ""
                level = 1
            elif fld_type == "separate":
                in_field_result = True
            elif fld_type == "end":
                if is_styleref and section_number:
                    num_parts = section_number.split(".")
                    parts.append(".".join(num_parts[:level]))
                elif is_ref and ref_label:
                    parts.append(ref_label)
                in_field = False
                in_field_result = False
                is_styleref = False
                is_ref = False
                ref_label = ""
        elif tag == f"{{{_W}}}instrText":
            instr = (node.text or "").strip()
            if re.match(r"STYLEREF", instr, re.IGNORECASE):
                is_styleref = True
                level = _styleref_level(instr)
            elif instr.upper().startswith("REF "):
                instr_parts = instr.split()
                if len(instr_parts) >= 2:
                    bm = instr_parts[1]
                    if bm in bm_to_label:
                        is_ref = True
                        ref_label = bm_to_label[bm]
        elif tag == f"{{{_W}}}noBreakHyphen":
            if not in_field or (in_field_result and not is_styleref and not is_ref):
                parts.append("\u2011")
        elif tag == f"{{{_W}}}t":
            if not in_field:
                parts.append(node.text or "")
            elif in_field_result and not is_styleref and not is_ref:
                parts.append(node.text or "")
            # STYLEREF/REF result: skip stale cached text

    return "".join(parts)


def _para_ref_resolved_text(para_e: Any, bm_to_label: dict[str, str]) -> str:
    """Like _para_text_elem but replaces REF field cached text with correct labels.

    Body paragraphs reference figure/table captions via REF fields whose cached
    result is the stale label (e.g. 'Рисунок Методы...-1').  We substitute
    the cached text with the correctly resolved label from *bm_to_label*
    (e.g. 'Рисунок 3‑1') so the AI context is clean.
    """
    parts: list[str] = []
    in_field = False
    in_field_result = False
    is_ref = False
    ref_label: str = ""

    for node in para_e.iter():
        tag = node.tag
        if tag == f"{{{_W}}}fldChar":
            fld_type = node.get(f"{{{_W}}}fldCharType", "")
            if fld_type == "begin":
                in_field = True
                in_field_result = False
                is_ref = False
                ref_label = ""
            elif fld_type == "separate":
                in_field_result = True
            elif fld_type == "end":
                if is_ref and ref_label:
                    parts.append(ref_label)
                in_field = False
                in_field_result = False
                is_ref = False
                ref_label = ""
        elif tag == f"{{{_W}}}instrText":
            instr = (node.text or "").strip()
            if instr.upper().startswith("REF "):
                instr_parts = instr.split()
                if len(instr_parts) >= 2:
                    bm = instr_parts[1]
                    if bm in bm_to_label:
                        is_ref = True
                        ref_label = bm_to_label[bm]
        elif tag == f"{{{_W}}}noBreakHyphen":
            if not in_field or (in_field_result and not is_ref):
                parts.append("\u2011")
        elif tag == f"{{{_W}}}t":
            if not in_field:
                parts.append(node.text or "")
            elif in_field_result and not is_ref:
                parts.append(node.text or "")
            # REF result: skip stale cached text (correct label appended at field end)

    return "".join(parts)


def _build_fig_table_dict_docx(
    doc: Document,
    sections: list[Section],
    para_to_section: dict[int, str] | None = None,
) -> list[FigTableEntry]:
    """Build a list of FigTableEntry from the whole document.

    Uses bookmark-based REF fields for docx plus a fallback text scan.
    The section_number for each mention is inferred from the surrounding text.

    *para_to_section* maps flat-paragraph index → section number.  When
    provided (built by the caller after auto-numbering) it is used directly;
    otherwise it is rebuilt here via _extract_heading_number (legacy path).
    """
    body = doc.element.body
    paras_elems = list(body.iter(f"{{{_W}}}p"))

    flat_paras = doc.paragraphs
    if para_to_section is None:
        para_to_section = {}
        current_sec = ""
        for idx, para in enumerate(flat_paras):
            if _heading_level(para) > 0:
                current_sec = _extract_heading_number(para.text.strip())
            para_to_section[idx] = current_sec

    # Build accurate per-element section list (replaces fraction-based heuristic)
    elem_section = _build_elem_section_list(paras_elems, flat_paras, para_to_section)

    # Collect captions
    entries: dict[str, FigTableEntry] = {}  # label → entry

    # Also build a bookmark→label map for REF-field lookup
    bm_to_label: dict[str, str] = {}

    for para_idx_e, para_e in enumerate(paras_elems):
        sec_num = elem_section[para_idx_e]
        text = _para_caption_text(para_e, sec_num).strip()
        m = _CAPTION_RE.match(text)
        if m:
            label = m.group(0).strip()
            entry = FigTableEntry(label=label, caption=text, section_number=sec_num)
            entries[label] = entry

            # Record bookmarks in this paragraph
            for bm in para_e.iter(f"{{{_W}}}bookmarkStart"):
                bm_name = bm.get(f"{{{_W}}}name", "")
                if bm_name:
                    bm_to_label[bm_name] = label

    # Renumber per-section before collecting mentions so that context text
    # resolved via REF fields already uses the renumbered labels.
    entries, bm_to_label = _renumber_entries_by_section(
        list(entries.keys()), entries, bm_to_label
    )

    # Collect mentions via REF fields
    for para_idx_e, para_e in enumerate(paras_elems):
        instr_texts = [
            node.text or ""
            for node in para_e.iter(f"{{{_W}}}instrText")
        ]
        for instr in instr_texts:
            instr = instr.strip()
            if not instr.upper().startswith("REF "):
                continue
            parts = instr.split()
            if len(parts) < 2:
                continue
            bm_name = parts[1]
            if bm_name not in bm_to_label:
                continue
            label = bm_to_label[bm_name]
            context_parts = []
            if para_idx_e > 0:
                context_parts.append(_para_ref_resolved_text(paras_elems[para_idx_e - 1], bm_to_label).strip())
            context_parts.append(_para_ref_resolved_text(para_e, bm_to_label).strip())
            if para_idx_e < len(paras_elems) - 1:
                context_parts.append(_para_ref_resolved_text(paras_elems[para_idx_e + 1], bm_to_label).strip())
            context = "\n".join(p for p in context_parts if p)
            sec_num = elem_section[para_idx_e]
            entries[label].mentions.append(FigTableMention(context=context, section_number=sec_num))


    # Fallback: text scan for mentions not covered by REF fields
    mentioned_via_ref: set[str] = set()
    for entry in entries.values():
        for m in entry.mentions:
            mentioned_via_ref.add(m.context[:50])  # rough dedup key

    for para_idx_e, para_e in enumerate(paras_elems):
        text = _para_text_elem(para_e).strip()
        for m in _MENTION_RE.finditer(text):
            raw = m.group(0)
            # Normalise to label key
            label_candidate = _normalise_mention(raw)
            if label_candidate not in entries:
                continue
            context_parts = []
            if para_idx_e > 0:
                context_parts.append(_para_text_elem(paras_elems[para_idx_e - 1]).strip())
            context_parts.append(text)
            if para_idx_e < len(paras_elems) - 1:
                context_parts.append(_para_text_elem(paras_elems[para_idx_e + 1]).strip())
            context = "\n".join(p for p in context_parts if p)
            if context[:50] in mentioned_via_ref:
                continue  # already captured by REF field
            sec_num = elem_section[para_idx_e]
            entries[label_candidate].mentions.append(
                FigTableMention(context=context, section_number=sec_num)
            )
            mentioned_via_ref.add(context[:50])

    return list(entries.values())


def _renumber_entries_by_section(
    label_order: list[str],
    entries: dict[str, "FigTableEntry"],
    bm_to_label: dict[str, str],
) -> tuple[dict[str, "FigTableEntry"], dict[str, str]]:
    """Renumber figure/table entries 1, 2, 3… within each (section, kind) group.

    Word's SEQ field uses a chapter-wide running counter so captions in section
    5.3 may have values 25–31 even though they are the 1st–7th figure of that
    section.  This function reassigns sequential per-section numbers so that
    the first figure in 5.3 gets label 'Рисунок 5.3‑1', the second
    'Рисунок 5.3‑2', etc.
    """
    sec_kind_count: dict[tuple[str, str], int] = {}
    remap: dict[str, str] = {}

    for old_label in label_order:
        entry = entries.get(old_label)
        if entry is None:
            continue
        m = re.match(r'^(Таблица|Рисунок)\b', old_label, re.IGNORECASE)
        if not m:
            remap[old_label] = old_label
            continue
        kind = m.group(1)
        key = (entry.section_number, kind)
        sec_kind_count[key] = sec_kind_count.get(key, 0) + 1
        seq = sec_kind_count[key]
        remap[old_label] = f"{kind} {entry.section_number}\u2011{seq}"

    new_entries: dict[str, FigTableEntry] = {}
    for old_label, entry in entries.items():
        new_label = remap.get(old_label, old_label)
        new_caption = entry.caption.replace(old_label, new_label, 1) if old_label in entry.caption else entry.caption
        new_entries[new_label] = FigTableEntry(
            label=new_label,
            caption=new_caption,
            section_number=entry.section_number,
            mentions=entry.mentions,
        )

    new_bm_to_label = {bm: remap.get(lbl, lbl) for bm, lbl in bm_to_label.items()}
    return new_entries, new_bm_to_label


def _normalise_mention(raw: str) -> str:
    """Convert a loose mention like 'таблицу 3.1-2' to 'Таблица 3.1-2'."""
    raw = raw.strip()
    if re.match(r"таблиц", raw, re.IGNORECASE):
        kind = "Таблица"
    else:
        kind = "Рисунок"
    num_m = re.search(r"[\d\.]+[-–—‑][\d\.]+|[\d\.]+", raw)
    if num_m:
        return f"{kind} {num_m.group(0)}"
    return raw


def _build_elem_section_list(
    paras_elems: list,
    flat_paras,
    para_to_section: dict[int, str],
) -> list[str]:
    """Return a section-number string for every element in *paras_elems*.

    Scans *paras_elems* in document order.  When an element is a top-level
    paragraph (found in *flat_paras* by object identity), the current section
    is updated from *para_to_section*.  Table-cell paragraphs (absent from
    *flat_paras*) inherit the last seen section.  This avoids the severe
    distortion caused by the old fraction-based heuristic when
    len(paras_elems) >> len(flat_paras).
    """
    flat_id_to_sec: dict[int, str] = {
        id(p._p): para_to_section.get(i, "")
        for i, p in enumerate(flat_paras)
    }
    result: list[str] = []
    current = ""
    for para_e in paras_elems:
        sec = flat_id_to_sec.get(id(para_e))
        if sec is not None:
            current = sec
        result.append(current)
    return result


# ---------------------------------------------------------------------------
# PDF parser
# ---------------------------------------------------------------------------

def _parse_pdf(file_path: str, range_spec: dict | None) -> DocData:
    pdf = fitz.open(file_path)
    sections: list[Section] = []

    for page_num in range(len(pdf)):
        page = pdf[page_num]
        blocks = page.get_text("blocks")
        page_text = "\n".join(
            b[4].strip() for b in blocks if b[6] == 0 and b[4].strip()
        )
        if not page_text:
            continue

        page_label = str(page_num + 1)
        sections.append(Section(
            number=page_label,
            title=f"Страница {page_label}",
            text=page_text,
            level=0,
        ))

    # Filter by range_spec
    if range_spec and range_spec.get("type") == "pages":
        items = range_spec.get("items", [])
        if items:
            sections = [s for s in sections if _page_in_range(int(s.number), items)]

    # Build figure/table dict for PDF (text scan only; already from filtered sections)
    fig_table_dict = _build_fig_table_dict_pdf(sections)

    if range_spec and range_spec.get("type") == "pages":
        items = range_spec.get("items", [])
        if items:
            fig_table_dict = [
                e for e in fig_table_dict
                if _page_in_range(int(e.section_number), items)
            ]

    return DocData(fmt="pdf", file_path=file_path,
                   sections=sections, fig_table_dict=fig_table_dict, raw_pdf=pdf)


def _build_fig_table_dict_pdf(sections: list[Section]) -> list[FigTableEntry]:
    entries: dict[str, FigTableEntry] = {}

    for sec in sections:
        lines = sec.text.splitlines()
        for line_idx, line in enumerate(lines):
            m = _CAPTION_RE.match(line.strip())
            if m:
                label = m.group(0).strip()
                if label not in entries:
                    entries[label] = FigTableEntry(
                        label=label, caption=line.strip(), section_number=sec.number
                    )

    for sec in sections:
        lines = sec.text.splitlines()
        for line_idx, line in enumerate(lines):
            for m in _MENTION_RE.finditer(line):
                raw = m.group(0)
                label_candidate = _normalise_mention(raw)
                if label_candidate not in entries:
                    continue
                ctx_start = max(0, line_idx - 1)
                ctx_end = min(len(lines), line_idx + 2)
                context = "\n".join(lines[ctx_start:ctx_end])
                entries[label_candidate].mentions.append(
                    FigTableMention(context=context, section_number=sec.number)
                )

    return list(entries.values())
